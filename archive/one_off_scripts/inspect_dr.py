from docx import Document

path = '/Users/finlayma84/Dropbox/FCUCC Bulletins/Epiphany/Epiphany B 2024/Music Epiphany B 2024.docx'
doc = Document(path)

print(f"========== {path.split('/')[-1]} ==========")
print(f"paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}\n")

for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"  {i:3}: {p.text}")