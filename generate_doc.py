"""
Generate or update a Music document for a given liturgical season.
"""

from datetime import date as date_class, timedelta, datetime
from pathlib import Path
import sqlite3
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dateutil.easter import easter

DB_PATH = 'repertoire.db'

# Where generated docs go. Override per-season via the doc_paths table.
DROPBOX_ROOT = Path.home() / 'Downloads' / 'Music Picker'

MONTHS_REV = ['', 'JAN', 'FEB', 'MAR', 'APR', 'MAY', 'JUN',
              'JUL', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC']
MONTHS = {m: i for i, m in enumerate(MONTHS_REV) if m}
SLOTS = ['Prelude', 'Min Music', 'Offering', 'Postlude']

SLOT_LABEL = {
    # Display labels for newly generated documents.
    # Database slot names remain the older internal names: Prelude, Min Music, Offering, Postlude.
    'Prelude': 'PRELUDE',
    'Min Music': 'MUSIC MINISTRY',
    'Offering': 'OFFERTORY',
    'Postlude': 'POSTLUDE',
}

SLOT_ALIASES = {
    # Labels/placeholders the updater should recognize in older or hand-edited docs.
    'Prelude': ['PRELUDE', 'PRELUDES', 'BELL CHOIR PRELUDE'],
    'Min Music': ['MUSIC MINISTRY', 'MIN MUSIC', 'CHORAL ANTHEM'],
    'Offering': ['OFFERTORY', 'OFFERING'],
    'Postlude': ['POSTLUDE'],
}

PASTOR_LINES = [
    'OPENING HYMN:',
    'CLOSING HYMN:',
]

DEFAULT_FONT = 'Gotham Book'
HEADING_FONT = 'Gotham Bold'
DATE_FONT = 'Gotham Medium'
BODY_SIZE = 12

COLOR_GREEN  = '3A7C22'
COLOR_PURPLE = '5B2C6F'
COLOR_RED    = 'A52A2A'
COLOR_GOLD   = 'B8860B'
COLOR_BLACK  = '000000'
COLOR_PLACEHOLDER = '777777'


def season_bounds(season, year):
    e = easter(year)
    if season == 'Lent':
        return (e - timedelta(days=46), e - timedelta(days=1))
    if season == 'Easter':
        # Eastertide runs through Pentecost Sunday.
        return (e, e + timedelta(days=49))
    if season == 'Pentecost':
        christmas = date_class(year, 12, 25)
        c_wd = christmas.weekday()
        days_back = (c_wd + 1) % 7 or 7
        sunday_before_christmas = christmas - timedelta(days=days_back)
        advent_start = sunday_before_christmas - timedelta(days=21)
        # Your Season after Pentecost document starts after Pentecost Sunday, i.e. Trinity Sunday.
        return (e + timedelta(days=56), advent_start - timedelta(days=1))
    if season == 'Advent/Christmas':
        christmas = date_class(year, 12, 25)
        c_wd = christmas.weekday()
        days_back = (c_wd + 1) % 7 or 7
        sunday_before_christmas = christmas - timedelta(days=days_back)
        advent_start = sunday_before_christmas - timedelta(days=21)
        jan6 = date_class(year + 1, 1, 6)
        j_wd = jan6.weekday()
        if j_wd == 6: ep = jan6
        elif j_wd < 3: ep = jan6 - timedelta(days=j_wd + 1)
        else: ep = jan6 + timedelta(days=6 - j_wd)
        return (advent_start, ep - timedelta(days=1))
    if season == 'Epiphany':
        jan6 = date_class(year, 1, 6)
        j_wd = jan6.weekday()
        if j_wd == 6: ep = jan6
        elif j_wd < 3: ep = jan6 - timedelta(days=j_wd + 1)
        else: ep = jan6 + timedelta(days=6 - j_wd)
        return (ep, e - timedelta(days=47))
    raise ValueError(f'Unknown season: {season}')


def lectionary_letter(year):
    return {1: 'A', 2: 'B', 0: 'C'}[year % 3]


def sundays_in_range(start, end):
    d = start
    while d.weekday() != 6:
        d += timedelta(days=1)
        if d > end:
            return []
    out = []
    while d <= end:
        out.append(d)
        d += timedelta(days=7)
    return out


def color_for_sunday(d, season):
    e = easter(d.year)
    pentecost = e + timedelta(days=49)
    palm = e - timedelta(days=7)
    if d == pentecost: return COLOR_RED
    if d == palm: return COLOR_RED
    if d == e: return COLOR_GOLD
    if d == pentecost + timedelta(days=7): return COLOR_GOLD
    if d.month == 10 and (d + timedelta(days=7)).month == 11: return COLOR_RED
    if season == 'Lent': return COLOR_PURPLE
    if season == 'Advent/Christmas':
        if d.month == 12 and d.day >= 24: return COLOR_GOLD
        if d.month == 1: return COLOR_GOLD
        return COLOR_PURPLE
    if season == 'Easter': return COLOR_GOLD
    return COLOR_GREEN


def get_pieces_for_date(conn, the_date):
    date_str = f'{MONTHS_REV[the_date.month]} {the_date.day:02d}'
    rows = conn.execute(
        "SELECT * FROM pieces WHERE calendar_year=? AND date=?",
        (the_date.year, date_str)
    ).fetchall()
    pieces = {s: None for s in SLOTS}
    hymns = []
    occasion = ''
    for r in rows:
        if r['slot'] == 'Hymn':
            hymns.append({'title': r['title'], 'hymn_no': r['hymn_no']})
        elif r['slot'] in SLOTS and r['chosen_by'] == 'Michael':
            pieces[r['slot']] = {
                'title': r['title'],
                'composer': r['composer'],
                'composer_dates': r['composer_dates'],
                'performer': r['performer'],
            }
        if r['occasion'] and not occasion:
            occasion = r['occasion']
    return {'pieces': pieces, 'hymns': hymns, 'occasion': occasion}


def add_run(p, text, font=DEFAULT_FONT, bold=False, italic=False,
            size_pt=BODY_SIZE, color_hex=None):
    r = p.add_run(text)
    r.font.name = font
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size_pt)
    if color_hex:
        r.font.color.rgb = RGBColor.from_string(color_hex)
    return r


def tight_paragraph(doc, keep_with_next=False, keep_together=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.keep_with_next = keep_with_next
    pf.keep_together = keep_together
    return p


def slot_placeholder(slot):
    """Curly-brace placeholder used in generated docs for empty Michael-owned slots."""
    return '{' + SLOT_LABEL[slot] + '}'


DATE_IN_TEXT_RE = re.compile(r'(?:\(?\s*b\.\s*\d{4}\s*\)?|\(?\s*\d{4}\s*[-–]\s*\d{4}\s*\)?)')


def text_contains(haystack, needle):
    return needle.casefold() in haystack.casefold()


def title_already_has_dates(title):
    return bool(DATE_IN_TEXT_RE.search(title))


def slot_line_text(slot, piece):
    """Produce the plain text for a slot line — placeholder if empty, formatted piece otherwise."""
    label = SLOT_LABEL[slot]
    if not piece or not piece.get('title'):
        return slot_placeholder(slot)
    title = (piece['title'] or '').strip()
    composer = (piece.get('composer') or '').strip()
    dates = (piece.get('composer_dates') or '').strip()

    # Many imported rows already have the composer and/or dates inside the title text.
    # Avoid doubling them in pastor-facing Word documents.
    if composer and not text_contains(title, composer):
        title = f"{title}, {composer}"
    if dates and not text_contains(title, dates.strip('()')) and not title_already_has_dates(title):
        title = f"{title} ({dates})"
    return f"{label}: {title}"


def clear_paragraph_runs(paragraph):
    """Remove all runs from a paragraph so it can be rewritten cleanly."""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def write_pastor_line(paragraph, label):
    """Write a pastor-owned line. The sync code intentionally ignores these."""
    add_run(paragraph, label, font=DEFAULT_FONT, bold=True)


def write_slot_paragraph(paragraph, slot, piece):
    """Write one Michael-owned music slot with nicer placeholder formatting.

    Empty slots remain easy for the sync code to recognize because the visible
    text is still exactly {PRELUDE}, {MUSIC MINISTRY}, etc. Filled slots use
    the newer pastor-facing labels while preserving plain, compact formatting.
    """
    if not piece or not piece.get('title'):
        add_run(
            paragraph,
            slot_placeholder(slot),
            font=DEFAULT_FONT,
            italic=True,
            color_hex=COLOR_PLACEHOLDER,
        )
        return

    label = SLOT_LABEL[slot]
    line = slot_line_text(slot, piece)
    prefix = f"{label}:"
    if line.upper().startswith(prefix):
        body = line[len(prefix):].strip()
        add_run(paragraph, f"{label}: ", font=DEFAULT_FONT)
        add_run(paragraph, body, font=DEFAULT_FONT)
    else:
        add_run(paragraph, line, font=DEFAULT_FONT)


def default_doc_path(season, year):
    """Where would this season's doc live by default?"""
    safe_season = season.replace('/', ' ')
    folder = DROPBOX_ROOT / f"{safe_season} Year {lectionary_letter(year)} {year}"
    filename = f"Music {safe_season} {year}.docx"
    return folder / filename


def generate_template(season, year):
    """Generate a new doc with placeholders. Returns (success, path_or_message)."""
    output_path = default_doc_path(season, year)

    if output_path.exists():
        return False, f"A doc already exists at {output_path}. Delete it manually to regenerate."

    output_path.parent.mkdir(parents=True, exist_ok=True)

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    start, end = season_bounds(season, year)
    sundays = sundays_in_range(start, end)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = DEFAULT_FONT
    style.font.size = Pt(BODY_SIZE)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    season_heading = {
        'Advent/Christmas': 'ADVENT/CHRISTMAS',
        'Epiphany': 'SEASON OF EPIPHANY',
        'Lent': 'SEASON OF LENT',
        'Easter': 'EASTER SEASON',
        'Pentecost': 'SEASON AFTER PENTECOST',
    }[season]
    title_color = {
        'Advent/Christmas': COLOR_PURPLE,
        'Epiphany': COLOR_GREEN,
        'Lent': COLOR_PURPLE,
        'Easter': COLOR_GOLD,
        'Pentecost': COLOR_GREEN,
    }[season]

    p = tight_paragraph(doc)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_run(p, f"{season_heading} | YEAR {lectionary_letter(year)} {year}",
            font=HEADING_FONT, bold=True, color_hex=title_color)
    tight_paragraph(doc)

    for sunday in sundays:
        data = get_pieces_for_date(conn, sunday)
        sunday_color = color_for_sunday(sunday, season)
        date_str = f'{MONTHS_REV[sunday.month]} {sunday.day:02d}'
        # Include the colon even when the occasion is blank so the pastor can type after it.
        header_text = f'{date_str}:' + (f" {data['occasion']}" if data['occasion'] else '')
        # Keep each Sunday block together when possible so a lone date heading
        # does not get stranded at the bottom of a page.
        p = tight_paragraph(doc, keep_with_next=True, keep_together=True)
        add_run(p, header_text, font=DATE_FONT, bold=True, color_hex=sunday_color)

        # Pastor-owned lines are intentionally left for manual entry.
        # The sync/update code never touches these.
        p = tight_paragraph(doc, keep_with_next=True, keep_together=True)
        write_pastor_line(p, 'OPENING HYMN:')

        for slot in ['Prelude', 'Min Music', 'Offering']:
            piece = data['pieces'][slot]
            p = tight_paragraph(doc, keep_with_next=True, keep_together=True)
            write_slot_paragraph(p, slot, piece)

        p = tight_paragraph(doc, keep_with_next=True, keep_together=True)
        write_pastor_line(p, 'CLOSING HYMN:')

        piece = data['pieces']['Postlude']
        p = tight_paragraph(doc, keep_with_next=False, keep_together=True)
        write_slot_paragraph(p, 'Postlude', piece)
        p.paragraph_format.space_after = Pt(8)

        tight_paragraph(doc)

    conn.close()
    doc.save(str(output_path))

    # Record the path in the database
    conn = sqlite3.connect(DB_PATH)
    now = datetime.now().isoformat(timespec='seconds')
    conn.execute('''
        INSERT OR REPLACE INTO doc_paths (season_year, doc_path, generated_at, last_synced_at)
        VALUES (?, ?, ?, ?)
    ''', (f'{season}-{year}', str(output_path), now, now))
    conn.commit()
    conn.close()

    return True, str(output_path)