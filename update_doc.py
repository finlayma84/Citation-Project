"""
Update an existing Music doc in place: replace slot lines with current database values.
Identifies slot lines by their label prefix (PRELUDE:, MIN MUSIC:, etc., or [PRELUDE: ...]).
Walks the doc paragraph-by-paragraph, using date headers to know which Sunday we're in.
"""

import re
import sqlite3
from datetime import date as date_class, datetime
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

from generate_doc import (
    DB_PATH, MONTHS, MONTHS_REV, SLOTS, SLOT_LABEL,
    DEFAULT_FONT, BODY_SIZE, get_pieces_for_date, slot_line_text,
)

# Detect a slot line: either "PRELUDE: ..." or "[PRELUDE: To be Chosen]"
SLOT_LINE_RE = re.compile(
    r'^(?:\[)?(' + '|'.join(re.escape(SLOT_LABEL[s]) for s in SLOTS) + r')[:\]]',
    re.I
)

# Detect a date header: "MAY 31 Trinity Sunday", "JUN 07 Second Sunday w/Communion | ..."
DATE_HEADER_RE = re.compile(r'^([A-Z]{3})\s+(\d{1,2})\b')


def slot_from_label(label_text):
    """Map a label string back to our canonical slot name."""
    up = label_text.upper().strip()
    for slot, label in SLOT_LABEL.items():
        if up == label:
            return slot
    return None


def rewrite_paragraph(paragraph, new_text):
    """Replace a paragraph's text while preserving the first run's formatting."""
    # Capture the format of the first run (font, size, color) so we can reapply it
    first_run_format = None
    if paragraph.runs:
        r = paragraph.runs[0]
        first_run_format = {
            'font': r.font.name,
            'size': r.font.size,
            'bold': r.bold,
            'italic': r.italic,
            'color': r.font.color.rgb if r.font.color and r.font.color.rgb else None,
        }
    # Clear all existing runs
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    # Add a single new run with the same formatting
    new_run = paragraph.add_run(new_text)
    if first_run_format:
        if first_run_format['font']:
            new_run.font.name = first_run_format['font']
        if first_run_format['size']:
            new_run.font.size = first_run_format['size']
        new_run.bold = first_run_format['bold']
        new_run.italic = first_run_format['italic']
        if first_run_format['color']:
            new_run.font.color.rgb = first_run_format['color']
    else:
        new_run.font.name = DEFAULT_FONT
        new_run.font.size = Pt(BODY_SIZE)


def update_doc_for_season(season, year):
    """Update the doc on disk for the given season+year. Returns (success, message)."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    row = conn.execute(
        "SELECT doc_path FROM doc_paths WHERE season_year=?",
        (f'{season}-{year}',)
    ).fetchone()
    if not row:
        conn.close()
        return False, f"No doc generated yet for {season} {year}"

    from pathlib import Path
    path = Path(row['doc_path'])
    if not path.exists():
        conn.close()
        return False, f"Doc no longer exists at {path}"

    doc = Document(str(path))

    current_date = None     # the date we're currently inside (date_class)
    current_year = year

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # Is this a date header?
        m = DATE_HEADER_RE.match(text)
        if m:
            mon_str = m.group(1)
            day = int(m.group(2))
            if mon_str in MONTHS:
                try:
                    current_date = date_class(current_year, MONTHS[mon_str], day)
                except ValueError:
                    current_date = None
            continue

        # Is this a slot line, and do we know which Sunday we're in?
        if current_date is None:
            continue
        m = SLOT_LINE_RE.match(text)
        if not m:
            continue
        slot = slot_from_label(m.group(1))
        if not slot:
            continue

        # Look up what the database says for this Sunday + slot
        data = get_pieces_for_date(conn, current_date)
        piece = data['pieces'][slot]
        new_text = slot_line_text(slot, piece)
        if new_text != text:
            rewrite_paragraph(paragraph, new_text)

    doc.save(str(path))

    # Record sync time
    conn.execute(
        "UPDATE doc_paths SET last_synced_at=? WHERE season_year=?",
        (datetime.now().isoformat(timespec='seconds'), f'{season}-{year}')
    )
    conn.commit()
    conn.close()

    return True, str(path)